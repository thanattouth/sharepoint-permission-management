from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SharePoint-Permission-Management-Customer-Delivery-TH.docx"

FONT = "Leelawadee UI"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "6B7280"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)
            for paragraph in row.cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def style_run(run, size: int = 11, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        style_run(run)
    return p


def add_body(doc: Document, text: str) -> None:
    p = add_paragraph(doc, text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        style_run(run)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        style_run(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    if level == 1:
        style_run(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        style_run(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        style_run(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    style_run(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    style_run(r2, size=10, color=INK)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [1.65, 4.85]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "หัวข้อ", bold=True, color=DARK_BLUE)
    set_cell_text(hdr[1], "รายละเอียด", bold=True, color=DARK_BLUE)
    set_cell_shading(hdr[0], LIGHT_BLUE)
    set_cell_shading(hdr[1], LIGHT_BLUE)
    for label, detail in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], detail)
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, color=DARK_BLUE)
        set_cell_shading(cell, LIGHT_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(INK)
    doc.add_paragraph()


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SharePoint Permission Management - Customer Tenant Delivery")
    style_run(run, size=8, color=MUTED)
    return doc


def build_doc() -> None:
    doc = setup_document()

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("คู่มือเตรียม Tenant ลูกค้าก่อน Deployment")
    style_run(run, size=22, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("SharePoint Permission Management")
    style_run(run, size=12, color=MUTED)

    add_callout(
        doc,
        "วัตถุประสงค์ของเอกสาร",
        "เอกสารนี้ใช้คุยกับลูกค้าก่อนส่งมอบระบบ เพื่อให้ลูกค้าเตรียม Microsoft Entra ID, SharePoint Online, App Roles, Graph consent, รายชื่อ site/library และค่า configuration ที่จำเป็นให้พร้อม",
    )

    add_heading(doc, "1. ภาพรวมการย้ายไป Tenant ลูกค้า")
    add_body(
        doc,
        "ระบบนี้เป็นเว็บที่ทำงานร่วมกับ Microsoft 365 tenant ของลูกค้าโดยตรง ผู้ใช้เข้าสู่ระบบด้วย Microsoft Entra ID จากนั้นระบบอ่าน SharePoint site/library ผ่าน Microsoft Graph และบันทึก audit log ลง SharePoint List",
    )
    add_bullets(
        doc,
        [
            "ลูกค้าต้องมี Microsoft 365 tenant พร้อม Microsoft Entra ID และ SharePoint Online",
            "ต้องมี App Registration ใหม่ใน tenant ลูกค้า",
            "ต้องมี admin consent สำหรับ Microsoft Graph permissions",
            "ต้องระบุ SharePoint sites และ document libraries ที่ระบบจะจัดการ",
            "ต้อง assign app roles ให้ผู้ใช้หรือกลุ่มผู้ใช้",
        ],
    )

    add_heading(doc, "2. Checklist สิ่งที่ลูกค้าต้องเตรียม")
    add_matrix_table(
        doc,
        ["หมวด", "สิ่งที่ต้องเตรียม", "ตัวอย่าง"],
        [
            ["Tenant", "Tenant ID และ primary domain", "contoso.com"],
            ["Admin", "ผู้ดูแล Entra และ SharePoint ที่สามารถตั้งค่า app/consent/site ได้", "Global Admin, SharePoint Admin"],
            ["Hosting", "URL สุดท้ายของเว็บ", "https://sp-access-contoso.azurewebsites.net"],
            ["Users/Groups", "รายชื่อ user หรือ group ที่จะใช้ระบบ", "SP Access Admins, SP Access Reviewers"],
            ["SharePoint", "รายชื่อ site และ library ที่ระบบต้องอ่าน/จัดการ", "contoso.sharepoint.com:/sites/Finance"],
            ["Governance", "ตำแหน่ง audit list และ reviewer scope list", "PermissionAuditLog, PermissionReviewScopes"],
        ],
        [1.25, 3.25, 2.0],
    )

    add_heading(doc, "3. ข้อมูลที่เราต้องขอจากลูกค้า")
    add_key_value_table(
        doc,
        [
            ("Tenant ID", "Directory tenant ID ของลูกค้า"),
            ("Client ID", "Application client ID หลังสร้าง App Registration"),
            ("Tenant domain", "Domain หลัก เช่น contoso.com"),
            ("Internal domains", "Domain ที่นับเป็น internal เช่น contoso.com, contoso.onmicrosoft.com"),
            ("App URL", "URL ที่ deploy เว็บจริง"),
            ("Target sites", "รายการ SharePoint site ที่ระบบจะโหลด"),
            ("Protected libraries", "ชื่อ library ที่ต้อง label เป็น protected เช่น Confidential, Secret"),
            ("Audit site/list", "ตำแหน่ง SharePoint List สำหรับเก็บ audit log"),
            ("Reviewer scope site/list", "ตำแหน่ง list สำหรับกำหนด owner scope ของ Reviewer"),
        ],
    )

    add_heading(doc, "4. Microsoft Entra App Registration")
    add_body(doc, "ให้ลูกค้าสร้าง App Registration ใหม่ใน tenant ของลูกค้า โดยแนะนำให้ใช้แบบ single-tenant เพื่อให้ใช้ได้เฉพาะคนในองค์กรลูกค้า")
    add_key_value_table(
        doc,
        [
            ("Name", "SharePoint Permission Management"),
            ("Supported account types", "Accounts in this organizational directory only"),
            ("Platform", "Single-page application (SPA)"),
            ("Redirect URI", "URL เว็บจริง เช่น https://sp-access-contoso.azurewebsites.net"),
            ("Local test URI", "http://localhost:3000 ถ้าต้องทดสอบ local"),
        ],
    )
    add_callout(
        doc,
        "ข้อควรระวัง",
        "ไม่ควรใช้ URL ของ My Apps launcher เป็น redirect URI เพราะเว็บนี้ใช้ MSAL ในตัวเอง ให้ My Apps เป็นเพียง tile สำหรับเปิดเว็บเท่านั้น",
    )

    add_heading(doc, "5. App Roles")
    add_body(doc, "ระบบใช้ app role จาก token เพื่อกำหนดสิทธิ์ของผู้ใช้ หลังสร้าง role แล้วต้อง assign user/group ผ่าน Enterprise Application")
    add_matrix_table(
        doc,
        ["Role value", "หน้าที่"],
        [
            ["Admin", "ดู site, review permission, ดู audit และ grant/update/remove direct permission ได้"],
            ["Reviewer", "อ่านอย่างเดียว ดู review report และ audit ได้"],
            ["GuestUser", "สิทธิ์ต่ำมาก ใช้สำหรับผู้ใช้ที่ไม่ควรเห็นข้อมูลหลัก"],
            ["InternalUser", "role เดิมสำหรับ read-only browsing แบบจำกัด"],
            ["ExecutiveUser", "role เดิมที่ behavior เหมือน Reviewer"],
        ],
        [1.6, 4.9],
    )
    add_body(doc, "ถ้าลูกค้าต้องการเหลือแค่ Admin, Reviewer และ GuestUser ให้ reassign ผู้ใช้ออกจาก InternalUser/ExecutiveUser ก่อน disable หรือ remove role เก่า")

    add_heading(doc, "6. Microsoft Graph Permissions")
    add_body(doc, "ลูกค้าต้องให้ admin consent กับ delegated Microsoft Graph permissions ต่อไปนี้")
    add_matrix_table(
        doc,
        ["Permission", "ใช้ทำอะไร"],
        [
            ["User.Read", "ใช้สำหรับ sign-in และอ่าน profile พื้นฐาน"],
            ["User.ReadBasic.All", "ใช้สำหรับค้นหา user และแสดงข้อมูลบุคคล"],
            ["Sites.Read.All", "อ่าน SharePoint site, library, file และ permission"],
            ["Sites.ReadWrite.All", "เขียน permission changes และ SharePoint list content"],
            ["Files.ReadWrite.All", "แก้ permission ของ drive item ผ่าน Microsoft Graph"],
        ],
        [2.0, 4.5],
    )
    add_callout(
        doc,
        "หลักการสิทธิ์",
        "ระบบใช้ delegated permissions ดังนั้นสิทธิ์จริงยังขึ้นกับผู้ใช้ที่ login ด้วย ผู้ใช้ที่เป็น Admin ใน app ควรมี SharePoint/Graph permission เพียงพอสำหรับการทำงานจริง",
    )

    add_heading(doc, "7. SharePoint ที่ต้องเตรียม")
    add_heading(doc, "7.1 Target Sites", level=2)
    add_body(doc, "ลูกค้าต้องแจ้ง site ที่ให้ระบบจัดการ โดยใช้รูปแบบ hostname:/sites/path")
    add_code_block(doc, "contoso.sharepoint.com:/sites/Finance,contoso.sharepoint.com:/sites/Legal")

    add_heading(doc, "7.2 Protected Libraries", level=2)
    add_body(doc, "ระบุชื่อ document library ที่ต้องนับเป็น protected เช่น Confidential หรือ Secret ชื่อต้องตรงกับ display name ของ library")
    add_code_block(doc, "Confidential,Secret")

    add_heading(doc, "7.3 Audit List", level=2)
    add_body(doc, "ระบบบันทึก audit log ลง SharePoint List เพื่อใช้เป็นหลักฐานการเปลี่ยน permission")
    add_code_block(doc, "PermissionAuditLog")
    add_bullets(
        doc,
        [
            "บันทึก login, refresh report, grant access, update role และ remove access",
            "เก็บ actor, target, site/library, role, approved request number, status และ error",
            "ระบบสามารถสร้าง list/columns ได้ถ้า Admin ที่ใช้งานมีสิทธิ์เพียงพอ",
        ],
    )

    add_heading(doc, "7.4 Reviewer Scope List", level=2)
    add_body(doc, "ใช้กำหนดว่า Reviewer หรือ Owner คนใดต้อง review site/library ใด")
    add_matrix_table(
        doc,
        ["Column", "คำอธิบาย"],
        [
            ["OwnerEmail", "email ของ owner หรือ reviewer ต้องมีอย่างน้อย column นี้"],
            ["OwnerName", "ชื่อที่แสดงใน dropdown"],
            ["OwnerRole", "OwnerRep, VP, EVP หรือ Reviewer"],
            ["SiteName / Hostname / Path", "ใช้จับคู่กับ SharePoint site"],
            ["LibraryName", "ใช้จำกัด scope เฉพาะ library"],
            ["Active", "Yes/No เพื่อเปิดปิด scope row"],
        ],
        [2.25, 4.25],
    )

    add_heading(doc, "8. Environment Variables สำหรับ Deployment")
    add_body(doc, "ค่าเหล่านี้ต้องเปลี่ยนเป็นของ tenant ลูกค้า และต้องตั้งก่อน build ถ้าใช้ GitHub Actions/Next.js")
    add_code_block(
        doc,
        """NEXT_PUBLIC_MSAL_CLIENT_ID=<customer-app-client-id>
NEXT_PUBLIC_MSAL_TENANT_ID=<customer-tenant-id>
NEXT_PUBLIC_APP_SESSION_MAX_MINUTES=480
NEXT_PUBLIC_TENANT_DOMAIN=contoso.com
NEXT_PUBLIC_INTERNAL_DOMAINS=contoso.com,contoso.onmicrosoft.com
NEXT_PUBLIC_TARGET_SITES=contoso.sharepoint.com:/sites/Finance
NEXT_PUBLIC_PROTECTED_LIBRARY_NAMES=Confidential,Secret
NEXT_PUBLIC_AUDIT_SITE=contoso.sharepoint.com:/sites/Governance
NEXT_PUBLIC_AUDIT_LIST_NAME=PermissionAuditLog
NEXT_PUBLIC_REVIEW_SCOPE_SITE=contoso.sharepoint.com:/sites/Governance
NEXT_PUBLIC_REVIEW_SCOPE_LIST_NAME=PermissionReviewScopes
NEXT_PUBLIC_REVIEW_SCAN_ITEM_LIMIT=2000""",
    )
    add_callout(
        doc,
        "สำคัญ",
        "ตัวแปร NEXT_PUBLIC_* ถูกฝังตอน next build ดังนั้นหาก deploy ผ่าน pipeline ต้องตั้งค่าเหล่านี้ในระบบ build ด้วย ไม่ใช่ตั้งเฉพาะ Azure App Settings หลัง build",
    )

    add_heading(doc, "9. My Apps")
    add_body(doc, "ถ้าลูกค้าต้องการให้ app ขึ้นใน Microsoft My Apps ให้ตั้งเป็น linked application tile")
    add_numbered(
        doc,
        [
            "ตั้ง Sign-on URL หรือ Homepage URL เป็น URL เว็บจริง",
            "ตั้ง Visible to users เป็น Yes",
            "Assign users/groups ให้ Enterprise Application",
            "ให้ tile เปิดเว็บโดยตรง ไม่ต้องให้ My Apps ทำ OIDC sign-in แทนเว็บ",
        ],
    )

    add_heading(doc, "10. UAT Checklist หลัง Deployment")
    add_matrix_table(
        doc,
        ["รายการทดสอบ", "ผลลัพธ์ที่คาดหวัง"],
        [
            ["Admin sign-in", "เข้า Admin workspace ได้"],
            ["Reviewer sign-in", "เข้า Reviewer report ได้และแก้ permission ไม่ได้"],
            ["Target sites", "site ที่ตั้งค่าไว้แสดงใน site picker"],
            ["Libraries", "document libraries แสดงถูกต้อง"],
            ["Permissions", "เห็น direct/inherited permissions"],
            ["Grant/Update/Remove", "Admin ทำได้เมื่อกรอก approved request number"],
            ["Audit log", "มี record ใน PermissionAuditLog"],
            ["Reviewer scope", "owner dropdown โหลดจาก PermissionReviewScopes"],
            ["External classification", "domain นอก internal แสดงเป็น external"],
            ["Session timeout", "ครบเวลาที่ตั้งไว้แล้วต้อง sign in ใหม่"],
        ],
        [2.4, 4.1],
    )

    add_heading(doc, "11. สิ่งที่ส่งมอบและความรับผิดชอบหลัง Go-live")
    add_bullets(
        doc,
        [
            "ลูกค้าดูแลการ assign app roles ให้ user/group",
            "ลูกค้าดูแล SharePoint sites/libraries ที่ต้องให้ระบบจัดการ",
            "ลูกค้าดูแล retention policy ของ audit log",
            "ลูกค้าดูแล Reviewer Scope List ให้ตรงกับ owner จริง",
            "ลูกค้าดูแล SharePoint sharing policy และ external collaboration policy",
            "Approval workflow อยู่นอกระบบนี้ ระบบจะเก็บ approved request number เพื่ออ้างอิงเท่านั้น",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
