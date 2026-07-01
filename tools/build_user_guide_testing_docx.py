from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/SP-Access-User-Guide-and-Testing.docx")


BLUE = "005EA8"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F6F8FA"
BORDER = "D9E2EC"
TEXT = "111827"
MUTED = "5F6B7A"
RED = "B91C1C"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])
                set_cell_border(row.cells[idx])
                row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_grid(table, widths: list[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    tbl = doc.add_table(rows=1, cols=len(headers))
    set_table_width(tbl, widths)
    set_table_grid(tbl, widths)
    header_cells = tbl.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(header_cells[idx], LIGHT_BLUE)
        p = header_cells[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        r.font.size = Pt(9.5)
    for row_data in rows:
        row = tbl.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string(TEXT)
            if value in {"Pass", "Fail", "Blocked", "Not tested"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return tbl


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_GRAY) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    set_table_width(tbl, [6.5])
    set_table_grid(tbl, [6.5])
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def screenshot_placeholder(doc: Document, caption: str, note: str = "วาง screenshot จริงในตำแหน่งนี้") -> None:
    tbl = doc.add_table(rows=1, cols=1)
    set_table_width(tbl, [6.5])
    set_table_grid(tbl, [6.5])
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "F9FBFD")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(f"Screenshot: {caption}\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(10)
    r2 = p.add_run(note)
    r2.font.color.rgb = RGBColor.from_string(MUTED)
    r2.font.size = Pt(9)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SP Access - User Guide and Scenario Testing")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SP Access")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("คู่มือการใช้งานและผลการทดสอบ Scenario")
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(TEXT)

    add_callout(
        doc,
        "วัตถุประสงค์ของเอกสาร",
        "เอกสารนี้ใช้เป็นคู่มือการเข้าใช้งาน การกำหนด permission และแบบบันทึกผลการทดสอบทุก scenario สำคัญของระบบ SharePoint Permission Management.",
    )

    table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ["Application URL", "https://spoapp.baht.net"],
            ["My Apps URL", "https://myapps.microsoft.com"],
            ["ระบบที่เกี่ยวข้อง", "Microsoft Entra ID, SharePoint Online, Microsoft Graph"],
            ["บทบาทหลัก", "Admin, Reviewer, Guest"],
            ["สถานะเอกสาร", "Draft สำหรับเติม screenshot และผลทดสอบจริง"],
        ],
        [1.8, 4.7],
    )

    doc.add_page_break()


def add_user_guide(doc: Document) -> None:
    doc.add_heading("1. ภาพรวมการใช้งาน", level=1)
    add_para(
        doc,
        "SP Access เป็นระบบสำหรับช่วย Admin ตรวจสอบและกำหนดสิทธิ์ SharePoint บนระดับ site, library, folder และ file โดยระบบใช้ Microsoft Graph และข้อมูล permission จาก SharePoint เป็นแหล่งข้อมูลหลัก.",
    )
    add_callout(
        doc,
        "ข้อควรรู้",
        "ระบบไม่ได้ bypass สิทธิ์ของ SharePoint ผู้ใช้ยังต้อง login ด้วย account ที่ได้รับ permission และ SharePoint จะประเมินสิทธิ์ล่าสุดทุกครั้งเมื่อเปิด link.",
    )

    doc.add_heading("2. ช่องทางเข้าใช้งาน", level=1)
    doc.add_heading("2.1 เข้าใช้งานผ่าน URL โดยตรง", level=2)
    add_numbered(
        doc,
        [
            "เปิด browser และเข้า https://spoapp.baht.net",
            "เลือก Sign in และ login ด้วยบัญชี Microsoft 365 ขององค์กร",
            "ตรวจสอบเมนูด้านซ้ายว่า role ที่ได้รับแสดงถูกต้อง เช่น Admin, Reviewer หรือ Guest",
        ],
    )
    screenshot_placeholder(doc, "หน้า login / หน้าแรกหลังเข้า https://spoapp.baht.net")

    doc.add_heading("2.2 เข้าใช้งานผ่าน My Apps", level=2)
    add_numbered(
        doc,
        [
            "เปิด https://myapps.microsoft.com",
            "ค้นหา application ชื่อ SP Access หรือชื่อที่องค์กรกำหนด",
            "กดเปิด application และตรวจสอบว่า redirect เข้าระบบได้สำเร็จ",
        ],
    )
    screenshot_placeholder(doc, "SP Access บน My Apps และหน้าหลัง launch สำเร็จ")

    doc.add_heading("3. บทบาทผู้ใช้งาน", level=1)
    table(
        doc,
        ["Role", "ใช้งานได้", "ข้อจำกัดหลัก"],
        [
            ["Admin", "จัดการ permission, grant/update/remove access, ดู audit", "ต้องมีสิทธิ์ SharePoint/Graph ที่เพียงพอ"],
            ["Reviewer", "ดู permission review และข้อมูล site ที่ตนมีสิทธิ์", "ไม่ควรแก้ permission"],
            ["Guest", "เปิดเฉพาะไฟล์หรือโฟลเดอร์ที่ได้รับสิทธิ์", "อาจมองไม่เห็น parent path ถ้าไม่ได้รับสิทธิ์"],
        ],
        [1.25, 3.0, 2.25],
    )

    doc.add_heading("4. ขั้นตอนกำหนด Permission", level=1)
    add_numbered(
        doc,
        [
            "เลือก site และ library ที่ต้องการจัดการ",
            "เลือก folder หรือ file เป้าหมาย",
            "เปิด panel Manage Access",
            "เลือกผู้รับสิทธิ์และ role เช่น viewer หรือ editor",
            "ยืนยันการ grant access",
            "copy item link จากผลลัพธ์หรือเมนูจุดสามจุด เพื่อส่งให้ผู้ใช้เป็น fallback",
            "ตรวจสอบ Audit Log ว่ามี action และ invite diagnostics ถูกต้อง",
        ],
    )
    screenshot_placeholder(doc, "เลือก site/library/folder/file")
    screenshot_placeholder(doc, "Confirm grant access และ copy item link")
    screenshot_placeholder(doc, "Audit Log หลัง grant access")

    doc.add_heading("5. การใช้ Link และ Email Invitation", level=1)
    add_bullets(
        doc,
        [
            "หากเปิด option ส่ง email ระบบจะเรียก SharePoint/Graph invitation flow ตาม permission ที่กำหนด",
            "หาก email ไม่ถึง โดยเฉพาะ external user ให้ใช้ direct item link เป็น fallback",
            "direct item link เป็น link เดิมได้ แต่ผลลัพธ์ขึ้นกับ permission ล่าสุดของ account ที่ login",
            "สำหรับไฟล์ที่อยู่ใต้ folder/library ที่ผู้ใช้ browse ไม่ได้ ควรส่ง direct file link ให้ผู้ใช้เปิดตรง",
        ],
    )
    add_callout(
        doc,
        "ข้อจำกัดเรื่อง parent permission",
        "ถ้า grant สิทธิ์เฉพาะไฟล์ เช่น Confidential > Folder1 > Readme.docx แต่ผู้ใช้ไม่มีสิทธิ์ที่ Confidential หรือ Folder1 ผู้ใช้อาจ browse จาก parent ไม่ได้ แต่ควรทดสอบการเปิดจาก direct file link แยกต่างหาก.",
        fill="FFF7ED",
    )


def add_test_plan(doc: Document) -> None:
    doc.add_heading("6. แผนการทดสอบ Scenario", level=1)
    add_para(
        doc,
        "ให้บันทึกผลทดสอบจริงพร้อม screenshot ทุก scenario โดยแยกระหว่าง inherited permission และ non-inherited permission เพื่อให้เห็น behavior ของ SharePoint ชัดเจน.",
    )

    doc.add_heading("6.1 Environment และบัญชีทดสอบ", level=2)
    table(
        doc,
        ["ประเภทบัญชี", "ตัวอย่าง", "ใช้ทดสอบ"],
        [
            ["User ในกลุ่ม", "user ภายใน tenant และอยู่ใน SharePoint group", "ยืนยันกรณี inherited/group permission"],
            ["User นอกกลุ่ม tenant เดียวกัน", "บัญชีภายในองค์กรแต่ไม่อยู่ในกลุ่ม", "ยืนยัน direct grant"],
            ["External guest", "บัญชีนอก tenant ที่มี guest status แล้ว", "ยืนยัน external access หลังถูก invite"],
            ["External non-guest", "บัญชีนอก tenant ที่ยังไม่มี guest status", "ยืนยัน invitation/onboarding flow"],
            ["VM/Desktop app", "เครื่องทดสอบแยก account", "ยืนยันเปิดไฟล์ผ่าน desktop app และ browser"],
        ],
        [1.7, 2.3, 2.5],
    )

    doc.add_heading("6.2 Folder Scenario", level=2)
    table(
        doc,
        ["Scenario", "Setup", "Expected Result", "Actual", "Evidence"],
        [
            ["Folder -> stop inherited group permission", "สร้าง folder ใหม่และ stop inheritance", "แสดง permission เฉพาะของ folder และ grant direct ได้", "", ""],
            ["Folder -> group permission still exist", "สร้าง folder ใหม่โดยยัง inherit group permission", "เห็น inherited/group permission และ direct grant เพิ่มได้", "", ""],
        ],
        [1.5, 1.8, 1.8, 0.7, 0.7],
    )

    doc.add_heading("6.3 Manage Access / Sharing Scenario", level=2)
    table(
        doc,
        ["Scenario", "Inheritance", "Target User", "Expected Result", "Actual", "Evidence"],
        [
            ["ให้ permission กับ user ในกลุ่ม", "Inherited", "Same group", "ผู้ใช้เปิดได้ตามสิทธิ์เดิม และ audit แสดง action ถูกต้อง", "", ""],
            ["ให้ permission กับ user นอกกลุ่ม แต่ tenant เดียวกัน", "Inherited", "Same tenant outside group", "ผู้ใช้เปิด direct link ได้ตาม role ที่ grant", "", ""],
            ["ให้ permission กับ user นอก tenant แต่มี guest status", "Inherited", "External guest", "ผู้ใช้เปิด direct link ได้หลัง login ด้วย guest account", "", ""],
            ["ให้ permission กับ user นอก tenant แต่ไม่มี guest status", "Inherited", "External non-guest", "เกิด invitation/onboarding และควรมี fallback direct link", "", ""],
            ["ให้ permission กับ user ในกลุ่ม", "Non-inherited", "Same group", "สิทธิ์บน item ถูกประเมินจาก direct/non-inherited permission", "", ""],
            ["ให้ permission กับ user นอกกลุ่ม แต่ tenant เดียวกัน", "Non-inherited", "Same tenant outside group", "ผู้ใช้เปิด item ได้เฉพาะ scope ที่ grant", "", ""],
            ["ให้ permission กับ user นอก tenant แต่มี guest status", "Non-inherited", "External guest", "ผู้ใช้เปิด item/folder ได้ตาม role", "", ""],
            ["ให้ permission กับ user นอก tenant แต่ไม่มี guest status", "Non-inherited", "External non-guest", "ตรวจ invite status, diagnostics และ fallback link", "", ""],
        ],
        [1.35, 1.0, 1.4, 1.85, 0.45, 0.45],
    )

    doc.add_heading("6.4 Validate Permission ที่ไฟล์", level=2)
    table(
        doc,
        ["Test Case", "Account", "Open Method", "Expected Result", "Actual", "Evidence"],
        [
            ["Viewer เปิดไฟล์", "Internal/External", "Browser", "เปิดอ่านได้ แต่แก้ไขไม่ได้", "", ""],
            ["Editor เปิดไฟล์", "Internal/External", "Browser", "เปิดและแก้ไขได้", "", ""],
            ["Viewer เปิดไฟล์ผ่าน desktop app", "Internal/External", "Desktop app on VM", "เปิดตามสิทธิ์และติด Rights Management ตามนโยบาย", "", ""],
            ["Direct file grant ใต้ parent ที่ไม่มีสิทธิ์", "Internal/External", "Direct file link", "เปิดไฟล์โดยตรงได้ถ้า SharePoint อนุญาต แต่ browse parent ไม่ได้", "", ""],
            ["Account ไม่มี permission แล้ว", "Internal/External", "Direct item link", "ถูก deny access", "", ""],
        ],
        [1.35, 1.2, 1.35, 2.0, 0.5, 0.5],
    )

    doc.add_heading("6.5 Audit Log Validation", level=2)
    table(
        doc,
        ["Check", "Expected Result", "Actual", "Evidence"],
        [
            ["Actor", "แสดงผู้ที่กด grant/update/remove จริง ไม่ใช่ชื่อคนเดียวเสมอ", "", ""],
            ["Target user", "แสดง email ผู้รับสิทธิ์ถูกต้อง", "", ""],
            ["Action", "GrantAccess / UpdateAccess / RemoveAccess ถูกต้อง", "", ""],
            ["Invite status", "แสดง Accepted / Partial / Failed / Unknown ตามผล Graph", "", ""],
            ["Share link", "มี link สำหรับ copy/open ได้หลังทำรายการ", "", ""],
        ],
        [1.7, 3.7, 0.55, 0.55],
    )

    doc.add_heading("7. Screenshot Checklist", level=1)
    table(
        doc,
        ["ลำดับ", "Screenshot ที่ต้องเก็บ", "สถานะ"],
        [
            ["1", "หน้าเข้าใช้งานผ่าน https://spoapp.baht.net", ""],
            ["2", "หน้าเปิดผ่าน My Apps", ""],
            ["3", "หน้าเลือก site/library", ""],
            ["4", "หน้า Manage Access ก่อน grant", ""],
            ["5", "หน้า Confirm grant access", ""],
            ["6", "ผลลัพธ์หลัง grant/update พร้อม copy link", ""],
            ["7", "Audit Log หลังทำรายการ", ""],
            ["8", "User ภายใน tenant เปิดไฟล์สำเร็จ", ""],
            ["9", "External guest เปิดไฟล์สำเร็จ", ""],
            ["10", "Desktop app/VM validation", ""],
        ],
        [0.65, 5.1, 0.75],
    )


def add_appendix(doc: Document) -> None:
    doc.add_heading("8. หมายเหตุและข้อจำกัด", level=1)
    add_bullets(
        doc,
        [
            "InviteDeliveryStatus = Accepted หมายถึง Graph/SharePoint รับคำสั่งสร้าง permission สำเร็จ ไม่ใช่หลักฐานว่า email ถูกส่งถึง inbox แล้ว",
            "External email delivery อาจขึ้นกับ tenant sharing policy, guest invitation flow, spam/junk, mailbox policy และข้อจำกัดของ SharePoint invitation",
            "หาก email delivery ไม่สม่ำเสมอ ให้ใช้ direct item link เป็น fallback หลัก และตรวจสอบ audit diagnostics ประกอบ",
            "Reviewer ควรเห็นเฉพาะข้อมูลที่ตนมีสิทธิ์ดู และไม่ควรต้องเลือก owner เองถ้า flow ถูกออกแบบให้ตรวจจาก account ที่ login",
        ],
    )

    add_callout(
        doc,
        "ผลลัพธ์สุดท้ายที่ต้องการ",
        "หลังทดสอบครบ เอกสารนี้ควรมี screenshot และค่า Actual Result ครบทุกแถว เพื่อใช้ส่งมอบให้ลูกค้าและเป็นหลักฐาน UAT/Delivery.",
    )


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_user_guide(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_test_plan(doc)
    add_appendix(doc)
    doc.save(OUT)


if __name__ == "__main__":
    main()
